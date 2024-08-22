import json
from docx import Document
from docx.shared import Inches
import os
import sys

def recreate_document(json_file, output_docx):
    # Load the JSON data
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Create a new Word document
    doc = Document()

    # Optional: Add the title as a heading (if desired)
    if data["title"]:
        doc.add_heading(data["title"], level=0)

    # Add sections to the document
    for section in data["sections"]:
        # Add section title as a heading
        doc.add_heading(section["section_title"], level=1)

        # Add content in the section
        for content in section["content"]:
            if content["type"] == "paragraph":
                doc.add_paragraph(content["text"])
            elif content["type"] == "list":
                for item in content["items"]:
                    doc.add_paragraph(item, style='List Bullet')
            elif content["type"] == "image":
                if os.path.exists(content["image_path"]):
                    doc.add_picture(content["image_path"], width=Inches(4.0))

    # Save the recreated document
    doc.save(output_docx)
    print(f"Document recreated and saved as {output_docx}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python recreate_doc.py <json_file> [<output_docx>]")
        sys.exit(1)

    json_file = sys.argv[1]

    # Use the second argument as the output file name, if provided
    output_docx = sys.argv[2] if len(sys.argv) > 2 else 'recreated_document.docx'

    recreate_document(json_file, output_docx)

if __name__ == "__main__":
    main()
